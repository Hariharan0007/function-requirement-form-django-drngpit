from django.shortcuts import render

import os
from docx import Document
from function_form import settings
import time
from datetime import datetime,date,timedelta
from docx2pdf import convert
from django.http import HttpResponse, Http404
from .models import admin_model,staff_model,function_model,venue_model,booking_model
from .serializers import admin_serializer,staff_serializer,function_serializer,venue_serializer,booking_serializer

# for mail connections
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


def login(request):
    if request.method=="POST":
        if(request.POST.get('login_as')=="Admin"):

            function_models = function_model.objects.all().order_by('-func_date')
            serializer_function = function_serializer(function_models,many = True)
            func_list = []
            for i in serializer_function.data:
                func_list.append(dict(i))
            if(len(func_list)==0):
                function = "none"
            else:
                function = func_list
            admin_mailid = request.POST.get('mail_id')
            admin_id = request.POST.get('staff_id')
            print(admin_mailid+" "+admin_id)

            admin_models = admin_model.objects.all()
            serializer_admin = admin_serializer(admin_models,many=True)
            admin_list = [] 
            for i in serializer_admin.data:
                for dit in dict(i):
                    print(dict(i)['email'])
                    print(dict(i)['admin_id'])
                    if dict(i)['email'] == admin_mailid:
                        if dict(i)['admin_id'] == admin_id:
                            print("Logged")
                            return render(request,'admin.html',{'function':function,'mailid':admin_mailid,})
                        else:
                            print("ELSe")
                            return render(request,'login_form.html',{'confirmation':"id_not_match"})
                    else:
                        print("else")
                        return render(request,'login_form.html',{'confirmation' : "failure"})
        else:
            admin_mailid = request.POST.get('mail_id')
            admin_id = request.POST.get('staff_id')
            print(admin_mailid+" "+admin_id)
            
            if(staff_model.objects.filter(email=admin_mailid)):
                if(staff_model.objects.filter(email=admin_mailid,staff_id=admin_id)):
                    return render(request,'function_form.html',{'mailid':admin_mailid,'view_form':'view'})
                else:
                     return render(request,'login_form.html',{'confirmation':"id_not_match"})
            else:
                return render(request,'login_form.html',{'confirmation' : "failure"})
    else:
        return render(request,'login_form.html',{})

# def update_function(request):
#         if request.method=="POST":
#                 mailid = request.POST.get('mailid')
#                 return render(request,'admin.html',{
#                         'mailid':mailid,
#                 })


def edit_form_test(request):
        if (request.method=="POST"):
                name = request.POST.get('name')
                profession = request.POST.get('pro')
                age = request.POST.get('age')
                return render(request,'edit_test.html',{
                        'name':name,
                        'pro':profession,
                        'age':age,
                })
        return render(request,'edit_test.html',{})

def edit_function(request):
        if request.method=="POST":
                mailid = request.POST.get('mailid')
                organizer_mail_id = request.POST.get('organizer_mail_id')
                func_name = request.POST.get('func_name')
                func_date = request.POST.get('func_date')
                print(mailid)
                print(organizer_mail_id)
                print(func_name)
                print(func_date)
                function_models = function_model.objects.filter(func_name = func_name,func_date=func_date,organizer_mail_id = organizer_mail_id)
                serialzer_function = function_serializer(function_models,many = True)
                func_list = []
                for i in serialzer_function.data:
                        func_list.append(dict(i))
                print(func_list)
                return render(request,'admin.html',{
                        'mailid':mailid,
                        'edit_function':func_list,
                        'view_form':'view',
                        })

def admin(request):
        if request.method=="POST":
                func_name = request.POST.get('func_name')
                venue = request.POST.get('func_venue')
                mailid = request.POST.get('mail_id')
                func_date = request.POST.get('func_date')
                start_time = request.POST.get('start_time')
                start_time = start_time[:5]
                end_time = request.POST.get('end_time')
                end_time = end_time[:5]
                org_mail = request.POST.get('org_mail')
                if request.POST.get('admin_action')=="approve":
                        booking_model.objects.filter(booking_date = func_date,starting_time = start_time,ending_time = end_time).update(status = "approved")
                        function_model.objects.filter(func_name = func_name,func_date = func_date,time_duration_start = start_time,time_duration_end = end_time).update(status = "approved")
                        body = "The function form for the booking of "+venue+" on "+func_date+" were successfully approved by admin."
                        send_mail(sender=org_mail,func_name="Nil",msg_subject="Function form - "+func_name+" - Approved",msg_body=body,msg_attachment="No",func_date=func_date)
                        print("Status approved")
                elif request.POST.get('admin_action')=="cancel":
                        booking_model.objects.filter(booking_date = func_date,starting_time = start_time,ending_time = end_time).update(status = "cancelled")
                        function_model.objects.filter(func_name = func_name,func_date = func_date,time_duration_start = start_time,time_duration_end = end_time).update(status = "cancelled")
                        body = "The function form for the booking of "+venue+" on "+func_date+" were cancelled by admin."
                        send_mail(sender=org_mail,func_name="Nil",msg_subject="Function form - "+func_name+" - Approved",msg_body=body,msg_attachment="No",func_date=func_date)
                        print("Status cancelled")

                func_list = function_list()
                return render(request,'admin.html',{
                        'mailid':mailid,
                        'function':func_list,
                })

def send_mail(sender,func_name,msg_subject,msg_body,msg_attachment,func_date):

        fromaddr = "hariharanp20ug0520@drngpit.ac.in"
        toaddr = sender

        # instance of MIMEMultipart
        msg = MIMEMultipart()

        # storing the senders email address
        msg['From'] = fromaddr

        # storing the receivers email address
        msg['To'] = toaddr

        # storing the subject
        msg['Subject'] = msg_subject

        # string to store the body of the mail
        body = msg_body

        # attach the body with the msg instance
        msg.attach(MIMEText(body, 'plain'))

        if(msg_attachment!="No"):
                # open the file to be sent
                filename = func_name + "_" + func_date + ".pdf"
                msg_attachment.replace("\\\\","//")
                attachment = open(msg_attachment, "rb")

                # instance of MIMEBase and named as p
                p = MIMEBase('application', 'octet-stream')

                # To change the payload into encoded form
                p.set_payload((attachment).read())

                # encode into base64
                encoders.encode_base64(p)

                p.add_header('Content-Disposition', "attachment; filename= %s" % filename)

                # attach the instance 'p' to instance 'msg'
                msg.attach(p)
                
                attachment.close()

        # creates SMTP session
        s = smtplib.SMTP('smtp.gmail.com', 587)

        # start TLS for security
        s.starttls()

        # Authentication
        s.login(fromaddr, "Drngpit@123")

        # Converts the Multipart msg into a string
        text = msg.as_string()

        # sending the mail
        s.sendmail(fromaddr, toaddr, text)

        # terminating the session
        s.quit()

        return True


def function_home(request):
        mail_id = request.POST.get('mail_id')
        status = request.POST.get('function_status')
        if(status=="waiting"):
                function_models = function_model.objects.filter(organizer_mail_id=mail_id,status="waiting")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'function_form.html',{
                        'mailid':mail_id,
                        'function':function,
                        'view_form':'no_view',
                })
        elif(status=="approved"):
                function_models = function_model.objects.filter(organizer_mail_id=mail_id,status="approved")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'function_form.html',{
                        'mailid':mail_id,
                        'function':function,
                        'view_form' : 'no_view',
                })
        elif(status=="cancelled"):
                function_models = function_model.objects.filter(organizer_mail_id=mail_id,status="cancelled")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'function_form.html',{
                        'mailid':mail_id,
                        'function':function,
                        'view_form' : 'no_view',
                })
        return render(request,'function_form.html',{
                'mailid':mail_id,
                'view_form':'view',
        })

def function_admin(request):
        mail_id = request.POST.get('mail_id')
        status = request.POST.get('function_status')
        staff_lst = request.POST.get('stafflist')
        if(staff_lst=='list'):
                staff_models = staff_model.objects.all()
                serializer_staff = staff_serializer(staff_models,many=True)
                staff_list = []
                for i in serializer_staff.data:
                        staff_list.append(dict(i))
                return render(request,'admin.html',{
                        'mailid':mail_id,
                        'staff':staff_list,
                        'staff_lst':'List',
                })
        if(status=="waiting"):
                function_models = function_model.objects.filter(status="waiting")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'admin.html',{
                        'mailid':mail_id,
                        'function':function,
                })
        elif(status=="approved"):
                function_models = function_model.objects.filter(status="approved")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'admin.html',{
                        'mailid':mail_id,
                        'function':function,
                })
        elif(status=="cancelled"):
                function_models = function_model.objects.filter(status="cancelled")
                serializer_func = function_serializer(function_models,many=True)
                func_list = []
                for i in serializer_func.data:
                        func_list.append(dict(i))
                if(len(func_list)==0):
                        function = 'none'
                else:
                        function = func_list
                return render(request,'admin.html',{
                        'mailid':mail_id,
                        'function':function,
                })
        function_models = function_model.objects.all()
        serializer_func = function_serializer(function_models,many=True)
        func_list = []
        for i in serializer_func.data:
                func_list.append(dict(i))
        if(len(func_list)==0):
                function = 'none'
        else:
                function = func_list
        return render(request,'admin.html',{
                'mailid':mail_id,
                'function':function,
        })



def function_list():
        function_models = function_model.objects.all()
        serializer_function = function_serializer(function_models,many = True)
        func_list = []
        for i in serializer_function.data:
                func_list.append(dict(i))
        return func_list


def check_availability(request):
    func_start_date = request.POST.get('func_start_date_check')
    func_end_date = request.POST.get('func_end_date_check')
    start_time = request.POST.get('time_duration_start_check')
    end_time = request.POST.get('time_duration_end_check')
    mail_id = request.POST.get('mail_id')
    internal_stud = int(request.POST.get('internal_stud_capacity'))
    external_stud = int(request.POST.get('external_stud_capacity'))
    capacity = internal_stud + external_stud
    func_days = 0

    print("Function start date ==>",func_start_date)
    print("Function end date ==>",func_end_date)

    func_st = str(func_start_date).split('-')
    func_end = str(func_end_date).split('-')

    month_list = ["January","Feburary","March","April","May","June","July","August","September","October","November","December"]
    month = month_list[(int(func_st[1])-1)]

    print(month)

    func_dates = []

    start_date = date(int(func_st[0]), int(func_st[1]), int(func_st[2]))
    end_date = date(int(func_end[0]), int(func_end[1]), int(func_end[2]))
    delta = timedelta(days=1)
    print("CHECK===>>>>",start_date <= end_date)

    while start_date <= end_date:  # func_st[1]
        func_dates.append(datetime.strftime(start_date, '%Y-%m-%d'))
        start_date += delta
        func_days+=1
           
    print("Func days : >>>",func_days)
    
    booking_models = booking_model.objects.filter(month=month)
    serializer_booking = booking_serializer(booking_models,many=True)

    date_collections = []
    temp_date = []
    for date_data in serializer_booking.data:
           temp_date.append(dict(date_data)['booking_date'])
           temp_date.append(dict(date_data)['func_days'])
           date_collections.append(temp_date)
           temp_date=[]

    # date_collections=[['2023-02-02',2],['2023-02-06',4],['2023-02-16',1]]

    bulk_date_collection = []
    temp_date=[]
    for dates in date_collections:
        for i in range(dates[1]):
            temp_date.append((datetime.strptime(dates[0], '%Y-%m-%d').date() + timedelta(i)).isoformat())
        #     print((datetime.strptime(dates[0], '%Y-%m-%d').date() + timedelta(i)).isoformat())
        bulk_date_collection.append(temp_date)
        temp_date=[]

    print('func_dates==>',func_dates)       
    print(bulk_date_collection)

    diff_venue_count = 0
    diff_booked_date = []
    for fn_date in bulk_date_collection:
        for f_date in fn_date:
            if f_date in func_dates:
                if f_date not in diff_booked_date:
                    diff_booked_date.append(fn_date[0])
                    diff_venue_count+=1
                #     print(fn_date[0])
                break
    print(diff_booked_date)
                    
    booked_venue = []
    for booked_date in diff_booked_date:
        booking_models = booking_model.objects.filter(booking_date=booked_date)
        serializer_booking = booking_serializer(booking_models,many=True)
        
        for data in serializer_booking.data:
            booked_venue.append(dict(data)['venue'])
    
    print("BV===>",booked_venue)
    venue_models = venue_model.objects.filter(capacity__gte=capacity)
    serializer_venue = venue_serializer(venue_models,many = True)
    venue_list = []
    
    for data in serializer_venue.data:
        venue_list.append(dict(data)['venue'])
    
    print("V==>",venue_list)

    free_venue = []

    for venue in venue_list:
        if venue not in booked_venue:
            free_venue.append(venue)

    print("FV==>",free_venue)

    # return render(request,'function_form.html',{
    #        'view_form':'view',
    #        "mailid":mail_id,
    # })

    return render(request,'function_form.html',{
        'venue':free_venue,
        'start_date':func_start_date,
        'end_date':func_end_date,
        'func_month':month,
        'start_time':start_time,
        'end_time':end_time,
        'func_days':func_days,
        'mailid':mail_id,
        'internal_stud':internal_stud,
        'external_stud':external_stud,
        'view_form':'view',
        })



def add_new(request):
        if request.method=="POST":
                staff_mail = request.POST.get('add_staff_mail_id')
                staff_code = request.POST.get('add_staff_code')
                add_staff_reg = staff_model(
                        email = staff_mail,
                        staff_id = staff_code,
                )
                add_staff_reg.save()
                body = "You were successfully added to function form user,Your code is " + staff_code +".Change this to your own custom password in your first login."
                send_mail(sender=staff_mail,func_name="Nil",msg_subject="Function form - Successfully added",msg_body=body,msg_attachment="No")
                mailid = request.POST.get('mail_id')
                print(mailid)
                func_list = function_list()
                return render(request,'admin.html',{
                        'mailid':mailid,
                        'function':func_list,
                        'add_new':"added",
                })

def change_password(request):
        if(request.method=='POST'):
                mailid = request.POST.get('staff_mail_id')
                old_password = request.POST.get('old_password')
                new_password = request.POST.get('confirm_new_password')
                if(staff_model.objects.filter(email = mailid)):
                        if(staff_model.objects.filter(email = mailid,staff_id = old_password)):
                                staff_model.objects.filter(email = mailid).update(staff_id = new_password)
                                body = "Your request for changing the password was successfull,you can now login with your new password \""+new_password+"\" ."
                                send_mail(sender=mailid,func_name="Nil",msg_subject="Function form - Password change successfull",msg_body=body,msg_attachment="No")
                                return render(request,'login_form.html',{
                                        'pass_change':'success',
                                })
                        else:
                                return render(request,'login_form.html',{
                                        'pass_change':'code_not_match',
                                })
                else:
                        return render(request,'login_form.html',{
                                'pass_change':'user_not_found',
                        })
        return render(request,'login_form.html',{})


def delete(request):
        func_name = request.POST.get('func_name')
        mail_id = request.POST.get('organizer_mail_id')
        func_date = request.POST.get('func_date')
        venue = request.POST.get('venue')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        function_model.objects.filter(func_name = func_name,organizer_mail_id = mail_id,func_date = func_date).delete()
        booking_model.objects.filter(booking_date = func_date,venue = venue,starting_time = start_time,ending_time = end_time).delete()
        doc_file = os.getcwd()+"\\static\\function_form\\function_documents\\"+func_name+"_"+func_date+".pdf"
        os.remove(doc_file)
        body = "The function form for the booking of "+venue+" on "+func_date+" were successfully deleted."
        send_mail(sender=mail_id,func_name=func_name,msg_subject="Function form - "+func_name+" - Booking deleted",msg_body=body,msg_attachment="No",func_date=func_date)
        function_models = function_model.objects.filter(organizer_mail_id = mail_id,status ='cancelled')
        serializer_function = function_serializer(function_models,many = True)
        func_list = []
        for i in serializer_function.data:
                func_list.append(dict(i))
        if(len(func_list)==0):
                function = 'none'
        else:
                function = func_list
        return render(request,'function_form.html',{
                'mailid':mail_id,
                'function':function,
                })


def download(request):
        func_name = request.POST.get('func_name')
        func_date = request.POST.get('func_date')
        file_path = os.getcwd()+"\\static\\function_form\\function_documents\\"+func_name+"_"+func_date+".pdf"
        if os.path.exists(file_path):
                with open(file_path, 'rb') as fh:
                        response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
                        response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                        return response
        raise Http404


def submit(request):
    print("IN Submit")
    if request.method=="POST":
        ac_arrangement = request.POST.get('ac_arrangement')
        audience = int(request.POST.get('audience'))
        chief_guest_name = request.POST.get('chief_guest_name')
        dept_name = request.POST.get('department')
        designation = request.POST.get('designation')
        time_duration_start = request.POST.get('time_duration_func_start')
        print("st time---->",time_duration_start)
        time_duration_end = request.POST.get('time_duration_func_end')
        print("end time---------->",time_duration_end)
        dias = int(request.POST.get('dias'))
        field_type = request.POST.get('field_type')
        func_date = request.POST.get('func_date')
        # newly added start
        func_start_date=func_date
        func_end_date = request.POST.get('function_end_date')
        func_month = request.POST.get('func_on_month')
        # newly added end
        print("date--------->",func_date)
        func_days = int(request.POST.get('func_days'))
        func_name = request.POST.get('func_name')
        func_students = request.POST.get('func_students')
        func_students_external = request.POST.get('func_students_external')
        func_students_external_class = request.POST.get('func_students_external_class')
        if (request.POST.get('func_students_dept')!=None):
                func_students_year_course = request.POST.get('func_students_year') + " - " + request.POST.get('func_students_dept') + " - " + request.POST.get('func_students_class')
        else:
               func_students_year_course = request.POST.get('func_students_year')


        if(request.POST.get('guest_house_persons')==''):
            guest_house = "No"
            guest_house_days = 0
            guest_house_from_date = "---"
            guest_house_persons = 0
            guest_house_to_date = "---"
        else:
            guest_house = "Yes"
            guest_house_days = int(request.POST.get('guest_house_days'))
            guest_house_from_date = request.POST.get('guest_house_from_date')
            guest_house_persons = int(request.POST.get('guest_house_persons'))
            guest_house_to_date = request.POST.get('guest_house_to_date')

        laptop = request.POST.get('laptop')
        lcd_projector = request.POST.get('lcd_projector')
        

        if(request.POST.get('tiffin')!=''):
                tiffin = int(request.POST.get('tiffin'))
                lunch_required_time = request.POST.get('lunch_required_time')
                normal_lunch = int(request.POST.get('normal_lunch'))
                spl_lunch_non_veg = int(request.POST.get('spl_lunch_non_veg'))
                spl_lunch_veg = int(request.POST.get('spl_lunch_veg'))
                lunch_exact_numbers = int(request.POST.get('lunch_exact_numbers'))
        else:
                tiffin = 0
                lunch_required_time = '00:00'
                normal_lunch = 0
                spl_lunch_non_veg = 0
                spl_lunch_veg = 0
                lunch_exact_numbers = 0
        

        memento = request.POST.get('memento')
        memento_quantity = int(request.POST.get('memento_quantity'))
        memento_worth = int(request.POST.get('memento_worth'))
        mic_arrangement = request.POST.get('mic_arrangement')
        mic_number = int(request.POST.get('mic_number'))
        organizer_contact = int(request.POST.get('organizer_contact'))
        organizer_mail_id = request.POST.get('mail_id')
        organizer_name = request.POST.get('organizer_name')
        payment_through = request.POST.get('payment_through')
        if(payment_through=="Others"):
               payment_through = request.POST.get('payment_through_others')
        photography = request.POST.get('photography')
        if(photography=="Yes"):
                photographer=request.POST.get('photographer')
        else:
                photographer= "---"
        reception_item_rec = request.POST.get('reception_item_rec')

        # refreshment_for_guest = request.POST.get('refreshment_for_guest')
        if(request.POST.get('refreshment_for_guest_number')==""):
            refreshment_for_guest = "No"
            refreshment_for_guest_number = 0
            refreshment_for_guest_time = "---"
            refreshment_guest_coffee = 0
            refreshment_guest_snacks = 0
            refreshment_guest_tea = 0
        else:
            refreshment_for_guest = "Yes"
            refreshment_for_guest_number = int(request.POST.get('refreshment_for_guest_number'))
            refreshment_for_guest_time = request.POST.get('refreshment_for_guest_time')
            refreshment_guest_coffee = int(request.POST.get('refreshment_guest_coffee'))
            refreshment_guest_snacks = int(request.POST.get('refreshment_guest_snacks'))
            refreshment_guest_tea = int(request.POST.get('refreshment_guest_tea'))

        # refreshment_for_student = request.POST.get('refreshment_for_student')

        if(request.POST.get('refreshment_for_student_number')==""):
            refreshment_for_student = "No"
            refreshment_for_student_number = 0
            refreshment_for_student_time = "---"
            refreshment_student_coffee = 0
            refreshment_student_snacks = 0
            refreshment_student_tea = 0
        else:
            refreshment_for_student = "Yes"
            refreshment_for_student_number = int(request.POST.get('refreshment_for_student_number'))
            refreshment_for_student_time = request.POST.get('refreshment_for_student_time')
            refreshment_student_coffee = int(request.POST.get('refreshment_student_coffee'))
            refreshment_student_snacks = int(request.POST.get('refreshment_student_snacks'))
            refreshment_student_tea = int(request.POST.get('refreshment_student_tea'))

        seating_arrangement_numbers = int(request.POST.get('seating_arrangement_numbers'))
        table_cloth_number = int(request.POST.get('table_cloth_number'))

        
        # time_duration_start = "10:30"
        # time_duration_end = "15:00"
        training_type = request.POST.get('train_type')
        
        
        # transport_req = request.POST.get('transport_req')
        if(request.POST.get('transport_req_date')==""):
            transport_req = "No"
            transport_req_date = "---"
            transport_pickup_time = "---"
            transport_drop_time = "---"
            transport_location = "---"
            transport_pickup_person_contact = 0
            transport_pickup_person_name = "---"
        else:
            transport_req = "Yes"
            transport_req_date = request.POST.get('transport_req_date')
            transport_pickup_time = request.POST.get('transport_pickup_time')
            transport_drop_time = request.POST.get('transport_drop_time')
            transport_location = request.POST.get('transport_location')
            transport_pickup_person_contact = int(request.POST.get('transport_pickup_person_contact'))
            transport_pickup_person_name = request.POST.get('transport_pickup_person_name')


        transport_students = request.POST.get('transport_students')
        if(request.POST.get('transport_students')=="Yes"):
                transport_students_number = int(request.POST.get('transport_students_number'))
                transport_students_stage = request.POST.get('transport_students_stage')
        else:
                transport_students_number = 0
                transport_students_stage = "---"

        type_of_mic = request.POST.get('type_of_mic')
        venue = request.POST.get('venue')

        remarks = "No remarks"

        # Approval assignments
        hod_status = "waiting"
        status = "waiting"
        principal_status = "waiting"

        named_tuple = time.localtime()
        print(named_tuple)
        date_time = time.strftime("%d/%m/%Y", named_tuple)
        time_stamp = date_time + " and " + datetime.today().strftime("%I:%M %p")

        doc_path = os.path.join(settings.BASE_DIR,'static/function_form/files/form.docx')
        print(doc_path)
        doc_path.replace("\\","//")
        print(doc_path)
        doc = Document(doc_path)

        for para in doc.paragraphs:
                for run in para.runs:
                        string = run.text.strip(' ')
                        string = string.strip("        ")
                        # print(string+">##")
                        temp = []
                        for i in string.split(' '):
                                if i!='':
                                        temp.append(i)
                        # print(temp)
                        temp_str = ""
                        for string in temp:
                                if string=="ac_ar":
                                        temp_str +=   ac_arrangement 
                                elif string=="audi":
                                        temp_str +=   str(audience)
                                elif string=="cgn":
                                        temp_str +=   chief_guest_name 
                                elif string=="dept_name":
                                        temp_str +=   dept_name 
                                elif string=="dsg":
                                        temp_str +=   designation 
                                elif string=="dias":
                                        temp_str +=    str(dias)
                                elif string=="field_t":
                                        temp_str +=   field_type 
                                elif string=="func_date":
                                        temp_str +=   func_date 
                                elif string=="func_days":
                                        temp_str +=    str(func_days)  + " days"
                                elif string=="func_name":
                                        temp_str +=   func_name 
                                elif string=="func_stud":
                                        temp_str +=    str(func_students) 
                                elif string=="fsyc":
                                        temp_str +=   func_students_year_course
                                elif string=="func_st":
                                        temp_str +=   func_students_external
                                elif string=="_u":
                                       temp_str += " "
                                elif string=="fsy":
                                        temp_str +=   func_students_external_class 
                                elif string=="gh":
                                        temp_str +=   guest_house 
                                elif string=="ghd":
                                        temp_str +=    str(guest_house_days)
                                elif string=="ghp":
                                        temp_str +=   str(guest_house_persons) 
                                elif string=="gh_f_dt":
                                        temp_str +=   guest_house_from_date 
                                elif string=="gh_t_dt":
                                        temp_str +=   guest_house_to_date 
                                elif string.strip(' ')=="lrt":
                                        temp_str +=  " " +   str(lunch_required_time )
                                elif string=="ltp":
                                        temp_str +=   laptop 
                                elif string=="lcd_ar":
                                        temp_str +=   lcd_projector 
                                elif string=="le_num":
                                        temp_str +=  " " +   str(lunch_exact_numbers )
                                elif string=="memento":
                                        temp_str +=   memento 
                                elif string=="m_w":
                                        temp_str += " " +   str(memento_worth)
                                elif string=="m_q":
                                        temp_str += " " +   str(memento_quantity)
                                elif string.strip(' ')=="m_ar":
                                        temp_str +=  "           "+  mic_arrangement +"        "
                                elif string.strip(' ')=="mic_n":
                                        temp_str +=    str(mic_number)
                                elif string.strip(' ')=="nm_lh":
                                        temp_str +=  " " +   str(normal_lunch)
                                elif string=="oc":
                                        temp_str +=    str(organizer_contact )
                                elif string=="org_name":
                                        temp_str +=   organizer_name 
                                elif string=="pymt":
                                        temp_str +=   payment_through 
                                elif string=="pic":
                                        temp_str +=   photography
                                elif string=="ptg":
                                        temp_str +=   photographer 
                                elif string=="reception_item_req":
                                        temp_str +=   reception_item_rec 
                                elif string=="r_g":
                                        temp_str +=    str(refreshment_for_guest)  + "   " + "Tea : " +   str(refreshment_guest_tea)  + "    " + "Coffee : " +   str(refreshment_guest_coffee)  + "      " + "Snacks : " +   str(refreshment_guest_snacks)
                                elif string=="-gtm-":
                                        temp_str +="                        : " +  refreshment_for_guest_time
                                elif string=="-rs-":
                                        temp_str +=    str(refreshment_for_student)  + "   " + "Tea : " +   str(refreshment_student_tea)  + "    " + "Coffee : " +   str(refreshment_student_coffee)  + "      " + "Snacks : " +   str(refreshment_student_snacks)
                                elif string=="--s-":
                                        temp_str +=  " " +  refreshment_for_student_time 
                                elif string=="san":
                                        temp_str +=    str(seating_arrangement_numbers )
                                elif string.strip(' ')=="l_n_v":
                                        temp_str +=  " " +   str(spl_lunch_non_veg )
                                elif string.strip(' ')=="l_v":
                                        temp_str +=  " "+   str(spl_lunch_veg )
                                elif string=="tcn":
                                        temp_str +=    str(table_cloth_number )
                                elif string=="tiffin":
                                        temp_str +=    str(tiffin)
                                elif string=="tm_d_e":
                                        temp_str +=   time_duration_end 
                                elif string=="tm_d_s":
                                        temp_str +=   time_duration_start 
                                elif string=="tn_t":
                                        temp_str +=   training_type 
                                elif string=="tr_pk_tm":
                                        temp_str +=   transport_pickup_time 
                                elif string=="tr_dp_tm":
                                        temp_str +=   transport_drop_time 
                                elif string=="tr_loc":
                                        temp_str +=   transport_location 
                                elif string=="tr_pc":
                                        temp_str +=    str(transport_pickup_person_contact)
                                elif string=="tr_nm":
                                        temp_str +=   transport_pickup_person_name  + " & "
                                elif string=="tr_rec":
                                        temp_str +=   transport_req 
                                elif string=="tr_rq_dt":
                                        temp_str +=   transport_req_date 
                                elif string=="trs":
                                        temp_str +=   str(transport_students) 
                                elif string=="trs_s":
                                        temp_str +=  transport_students_stage  + " & " +   str(transport_students_number )
                                elif string.strip(' ')=="mic_type":
                                        temp_str +=   type_of_mic +" & "
                                elif string=="venue":
                                        temp_str +=   venue
                                elif string=="dt":
                                        temp_str +=   " " + time_stamp 
                                else:
                                        pass
                        if temp_str!="":
                                run.text = temp_str
        temp_file = doc.save(func_name+".docx")
        temp_file_path = os.getcwd()+"\\"+func_name+".docx"
        convert(temp_file_path,os.getcwd()+"\\static\\function_form\\function_documents\\"+func_name+"_"+func_date+".pdf")
        os.remove(temp_file_path)
        doc_file = os.getcwd()+"\\static\\function_form\\function_documents\\"+func_name+"_"+func_date+".pdf"
        
        reg = function_model(
            ac_arrangement = ac_arrangement,
            audience = audience,
            chief_guest_name = chief_guest_name,
            dept_name = dept_name,
            designation = designation,
            dias = dias,
            field_type = field_type,
            func_date = func_date,
            func_end_date = func_end_date,
            func_days = func_days,
            func_name = func_name,
            func_students = func_students,
            func_students_external = func_students_external,
            func_students_external_class = func_students_external_class,
            func_students_year_course = func_students_year_course,
            guest_house = guest_house,
            guest_house_days = guest_house_days,
            guest_house_from_date = guest_house_from_date,
            guest_house_persons = guest_house_persons,
            guest_house_to_date = guest_house_to_date,
            laptop = laptop,
            lcd_projector = lcd_projector,
            lunch_exact_numbers = lunch_exact_numbers,
            lunch_required_time = lunch_required_time,
            memento = memento,
            memento_quantity = memento_quantity,
            memento_worth = memento_worth,
            mic_arrangement = mic_arrangement,
            mic_number = mic_number,
            normal_lunch = normal_lunch,
            organizer_contact = organizer_contact,
            organizer_mail_id = organizer_mail_id,
            organizer_name = organizer_name,
            payment_through = payment_through,
            photography = photography,
            photographer = photographer,
            reception_item_rec = reception_item_rec,
            refreshment_for_guest = refreshment_for_guest,
            refreshment_for_guest_number = refreshment_for_guest_number,
            refreshment_for_guest_time = refreshment_for_guest_time,
            refreshment_for_student = refreshment_for_student,
            refreshment_for_student_number = refreshment_for_student_number,
            refreshment_for_student_time = refreshment_for_student_time,
            refreshment_guest_coffee = refreshment_guest_coffee,
            refreshment_guest_snacks = refreshment_guest_snacks,
            refreshment_guest_tea = refreshment_guest_tea,
            refreshment_student_coffee = refreshment_student_coffee,
            refreshment_student_snacks = refreshment_student_snacks,
            refreshment_student_tea = refreshment_student_tea,
            seating_arrangement_numbers = seating_arrangement_numbers,
            spl_lunch_non_veg = spl_lunch_non_veg,
            spl_lunch_veg = spl_lunch_veg,
            table_cloth_number = table_cloth_number,
            tiffin = tiffin,
            time_duration_end = time_duration_end,
            time_duration_start = time_duration_start,
            training_type = training_type,
            transport_drop_time = transport_drop_time,
            transport_location = transport_location,
            transport_pickup_person_contact = transport_pickup_person_contact,
            transport_pickup_person_name = transport_pickup_person_name,
            transport_pickup_time = transport_pickup_time,
            transport_req = transport_req,
            transport_req_date = transport_req_date,
            transport_students = transport_students,
            transport_students_number = transport_students_number,
            transport_students_stage = transport_students_stage,
            type_of_mic = type_of_mic,
            venue = venue,
            time_stamp = time_stamp,
            remarks = remarks,
            hod_status = hod_status,
            status = status,
            principal_status = principal_status,
            document_pdf = doc_file,
        )

        preview_location = func_name + "_" + func_date + ".pdf" 

        # print("Form action ====>>>",request.POST.get('form_action'))
        # doc_path = os.path.join(settings.BASE_DIR,'static/function_form/files/form.docx')
        # print(doc_path)
        # doc_path.replace("\\","//")
        # print(doc_path)
        # doc = Document(doc_path)

        if(request.POST.get('function_edit')=='edit_func'):
                admin_mail = request.POST.get('mail_id')
                function_model.objects.filter(func_name = func_name,organizer_mail_id = organizer_mail_id,func_date = func_date).delete()
                msg_sub = "Function form - " + func_name + " - Booking Updated !!! "
                msg_body = "The function form for booking "+venue+" on "+func_date+" were successfully updated."
                mail_status = "not_sent"
                if(send_mail(organizer_mail_id,func_name,msg_sub,msg_body,doc_file,func_date)):
                        mail_status = "sent"
                reg.save()
                # function_models = function_model.objects.all()
                # serializer_function = function_serializer(function_models,many = True)
                # func_list = []
                # for i in serializer_function.data:
                #         func_list.append(dict(i))
                # if(len(func_list)==0):
                #         function = "none"
                # else:
                #         function = func_list
                print("Function Updated to database !!!!!!! ",reg)
                return render(request,'admin.html',{
                       'mailid':admin_mail,
                       'document':preview_location,
                       'mail_status':mail_status,
                })


                
        reg.save()
        print("Function saved to database !!!!!!! ",reg)

        b_reg = booking_model(
        venue = venue,
        booking_date = func_date,
        booking_end_date = func_end_date,
        starting_time = time_duration_start,
        ending_time = time_duration_end,
        month = func_month,
        func_days = func_days,
        status = "waiting",
        )

        b_reg.save()

        print("Bookings Successfull !!!")

        # doc_file = re.sub(r"\s+","%20",doc_file)
        # print(doc_file)

        function_models = function_model.objects.filter(organizer_mail_id = organizer_mail_id,func_name = func_name)
        serializer_function = function_serializer(function_models,many = True)
        func_list = []
        for i in serializer_function.data:
                func_list.append(dict(i))
        msg_sub = "Function form - " + func_name + " - Booking submitted"
        msg_body = "The function form for booking "+venue+" on "+func_date+" were successfully submitted for approval."
        mail_status = "not_sent"
        if(send_mail(organizer_mail_id,func_name,msg_sub,msg_body,doc_file,func_date)):
                mail_status = "sent"



        return render(request,'function_form.html',{
                'mailid':organizer_mail_id,
                'document':preview_location,
                'function':func_list,
                'mail_status':mail_status,
                })
    else:
        pass